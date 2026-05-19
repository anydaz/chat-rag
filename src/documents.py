"""Document loader for handling multiple document types."""

from pathlib import Path
from typing import List, Tuple
from pypdf import PdfReader
from docx import Document


class DocumentLoader:
    """Generic document loader for different file types."""
    
    def __init__(self, directory: str = "data/pdfs"):
        self.directory = Path(directory)
    
    def load_all_documents(self) -> List[Tuple[str, dict]]:
        """Load all supported documents from the directory.
        
        Returns:
            List of tuples: (text, metadata)
            where metadata contains source, filename, and file type info
        """
        if not self.directory.exists():
            print(f"❌ Directory does not exist: {self.directory}")
            return []
        
        documents = []
        files = list(self.directory.glob("*"))
        
        print(f"Found {len(files)} file(s)")
        
        for file_path in files:
            if file_path.is_file():
                text, metadata = self.load_document(file_path)
                if text:
                    documents.append((text, metadata))
        
        return documents
    
    def load_document(self, file_path: Path) -> Tuple[str, dict]:
        """Load a single document based on its file type.
        
        Returns:
            Tuple of (text, metadata)
        """
        suffix = file_path.suffix.lower()
        metadata = {
            "filename": file_path.name,
            "file_type": suffix,
            "source_path": str(file_path),
            "directory": self.directory.name,
        }
        
        if suffix == ".pdf":
            text = self._load_pdf(file_path)
            metadata["document_type"] = "pdf"
        elif suffix in [".txt", ".md"]:
            text = self._load_text(file_path)
            metadata["document_type"] = "text"
        elif suffix == ".docx":
            text = self._load_docx(file_path)
            metadata["document_type"] = "word"
        else:
            print(f"  ⚠️  Unsupported file type: {suffix}")
            text = ""
            metadata["document_type"] = "unsupported"
        
        return text, metadata
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load text from PDF file."""
        try:
            print(f"  Reading PDF: {file_path.name}")
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            print(f"    Pages: {num_pages}")
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        except Exception as e:
            print(f"  ❌ Error reading PDF {file_path.name}: {e}")
            return ""
    
    def _load_text(self, file_path: Path) -> str:
        """Load text from plain text file."""
        try:
            print(f"  Reading text file: {file_path.name}")
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            print(f"    Size: {len(text)} characters")
            return text
        except Exception as e:
            print(f"  ❌ Error reading {file_path.name}: {e}")
            return ""
    
    def _load_docx(self, file_path: Path) -> str:
        """Load text from Word document."""
        try:
            print(f"  Reading DOCX: {file_path.name}")
            doc = Document(file_path)
            
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            print(f"    Size: {len(text)} characters")
            return text
        except Exception as e:
            print(f"  ❌ Error reading {file_path.name}: {e}")
            return ""
